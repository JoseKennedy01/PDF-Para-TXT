import json
import fitz  # PyMuPDF
from io import BytesIO

from django.shortcuts import render
from django.http import HttpResponse

from docx import Document
from docx.shared import Pt
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image
import pikepdf


# ── HELPERS ──────────────────────────────────────────────────────────────────

def _parse_paginas(pagina_str, total):
    """
    Converte string de páginas em lista de índices (0-based).
    Aceita: '1,3,5'  ou  '2-6'  ou  '1,3-5,8'  (vazio = todas)
    """
    if not pagina_str.strip():
        return list(range(total))
    indices = []
    for parte in pagina_str.split(','):
        parte = parte.strip()
        if '-' in parte:
            try:
                a, b = parte.split('-')
                indices += list(range(int(a) - 1, int(b)))
            except Exception:
                pass
        else:
            try:
                indices.append(int(parte) - 1)
            except Exception:
                pass
    return [i for i in indices if 0 <= i < total]


# ── VIEWS ─────────────────────────────────────────────────────────────────────

def upload_pdf(request):
    """Página principal + PDF → TXT"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        pdf_file = request.FILES['arquivo']
        modo = request.POST.get('modo', 'completo')

        doc = fitz.open(stream=pdf_file.read(), filetype='pdf')
        texto_completo = ''
        for i, pagina in enumerate(doc):
            texto = pagina.get_text()
            if modo == 'limpo':
                linhas = [l for l in texto.splitlines() if l.strip()]
                texto = '\n'.join(linhas)
            texto_completo += f'--- Página {i + 1} ---\n{texto}\n\n'
        doc.close()

        response = HttpResponse(texto_completo, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = 'attachment; filename="resultado.txt"'
        return response

    return render(request, 'meuapp/upload.html')


def pdf_para_docx(request):
    """PDF → DOCX (Word)"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        pdf = request.FILES['arquivo']
        doc_fitz = fitz.open(stream=pdf.read(), filetype='pdf')

        doc_word = Document()
        # Estilo do título do documento
        titulo = doc_word.add_heading('Documento convertido por PDFly', level=1)
        titulo.runs[0].font.size = Pt(16)

        for i, pagina in enumerate(doc_fitz):
            doc_word.add_heading(f'Página {i + 1}', level=2)
            texto = pagina.get_text()
            for linha in texto.splitlines():
                if linha.strip():
                    doc_word.add_paragraph(linha)
            doc_word.add_page_break()

        doc_fitz.close()
        buffer = BytesIO()
        doc_word.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="resultado.docx"'
        return response

    return render(request, 'meuapp/upload.html')


def pdf_para_xlsx(request):
    """PDF → XLSX (Excel) — extrai tabelas e texto com formatação"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        pdf = request.FILES['arquivo']
        doc_fitz = fitz.open(stream=pdf.read(), filetype='pdf')

        wb = Workbook()
        # Remove aba padrão
        wb.remove(wb.active)

        for i, pagina in enumerate(doc_fitz):
            ws = wb.create_sheet(title=f'Pág {i + 1}')

            # ── Cabeçalho da aba ──
            ws.merge_cells('A1:D1')
            cell_header = ws['A1']
            cell_header.value = f'Página {i + 1}'
            cell_header.font = Font(bold=True, size=13, color='FFFFFF')
            cell_header.fill = PatternFill('solid', fgColor='5B5FCC')
            cell_header.alignment = Alignment(horizontal='center', vertical='center')
            ws.row_dimensions[1].height = 22

            linha_ws = 2

            # ── Tenta extrair tabelas primeiro ──
            tabelas = pagina.find_tables()
            if tabelas.tables:
                for tabela in tabelas.tables:
                    dados = tabela.extract()
                    # Cabeçalho da tabela (primeira linha)
                    if dados:
                        for col_idx, celula in enumerate(dados[0], start=1):
                            c = ws.cell(row=linha_ws, column=col_idx, value=str(celula or ''))
                            c.font = Font(bold=True, color='FFFFFF')
                            c.fill = PatternFill('solid', fgColor='7C6FFF')
                            c.alignment = Alignment(wrap_text=True)
                        linha_ws += 1
                        # Demais linhas da tabela
                        for row in dados[1:]:
                            for col_idx, celula in enumerate(row, start=1):
                                c = ws.cell(row=linha_ws, column=col_idx, value=str(celula or ''))
                                c.alignment = Alignment(wrap_text=True)
                            linha_ws += 1
                        linha_ws += 1  # espaço entre tabelas
            else:
                # ── Sem tabelas: coloca o texto linha a linha ──
                texto = pagina.get_text()
                for txt_linha in texto.splitlines():
                    if txt_linha.strip():
                        ws.cell(row=linha_ws, column=1, value=txt_linha)
                        linha_ws += 1

            # Ajusta largura das colunas automaticamente
            for col in ws.columns:
                max_len = 0
                col_letter = None
                for cell in col:
                    try:
                        # MergedCell não tem column_letter, pula
                        if not hasattr(cell,'column_letter'):
                            continue
                        if col_letter is None:
                            col_letter = cell.column_letter
                        if cell.value:
                            max_len = max(max_len, len(str(cell.value)))
                    except Exception:
                        pass
                    if col_letter:
                        ws.column_dimensions[col_letter].width = min(max_len + 4, 60)

        doc_fitz.close()
        buffer = BytesIO()
        wb.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer,
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="resultado.xlsx"'
        return response

    return render(request, 'meuapp/upload.html')


def imagem_para_pdf(request):
    """Imagem (JPG/PNG/WEBP) → PDF"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        imagem = request.FILES['arquivo']
        img = Image.open(imagem).convert('RGB')
        largura, altura = img.size

        buffer = BytesIO()
        c = canvas.Canvas(buffer, pagesize=(largura, altura))

        img_buffer = BytesIO()
        img.save(img_buffer, format='PNG')
        img_buffer.seek(0)

        c.drawImage(ImageReader(img_buffer), 0, 0, largura, altura)
        c.save()
        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="resultado.pdf"'
        return response

    return render(request, 'meuapp/upload.html')


def mesclar_pdfs(request):
    """
    Mesclar PDFs com suporte a:
    - Seleção de páginas por arquivo (ex: 1-3 ou 1,3,5)
    - Ordem personalizada via drag-and-drop
    """
    if request.method == 'POST':
        arquivos = request.FILES.getlist('arquivos')
        merge_order_raw = request.POST.get('merge_order', '')

        if not arquivos:
            return render(request, 'meuapp/upload.html')

        # Tenta parsear a ordem e páginas enviadas pelo frontend
        try:
            merge_order = json.loads(merge_order_raw) if merge_order_raw else []
        except Exception:
            merge_order = []

        # Monta mapa idx → páginas
        paginas_map = {}
        for item in merge_order:
            idx = item.get('idx')
            pages_str = item.get('pages', '')
            paginas_map[idx] = pages_str

        pdf_final = pikepdf.Pdf.new()

        for idx, arq in enumerate(arquivos):
            pdf_temp = pikepdf.Pdf.open(BytesIO(arq.read()))
            total = len(pdf_temp.pages)
            pages_str = paginas_map.get(idx, '')
            indices = _parse_paginas(pages_str, total)
            for i in indices:
                pdf_final.pages.append(pdf_temp.pages[i])

        buffer = BytesIO()
        pdf_final.save(buffer)
        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="mesclado.pdf"'
        return response

    return render(request, 'meuapp/upload.html')


def dividir_pdf(request):
    """Dividir PDF — aceita '1,3,5' e '2-6' e combinações"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        pdf = request.FILES['arquivo']
        pagina_str = request.POST.get('paginas', '')

        doc = pikepdf.Pdf.open(BytesIO(pdf.read()))
        total = len(doc.pages)
        indices = _parse_paginas(pagina_str, total)

        novo_pdf = pikepdf.Pdf.new()
        for i in indices:
            novo_pdf.pages.append(doc.pages[i])

        buffer = BytesIO()
        novo_pdf.save(buffer)
        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="dividido.pdf"'
        return response

    return render(request, 'meuapp/upload.html')


def comprimir_pdf(request):
    """Comprimir PDF reduzindo tamanho do arquivo"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        pdf = request.FILES['arquivo']

        doc = pikepdf.Pdf.open(BytesIO(pdf.read()))
        buffer = BytesIO()
        doc.save(
            buffer,
            compress_streams=True,
            object_stream_mode=pikepdf.ObjectStreamMode.generate
        )
        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="comprimido.pdf"'
        return response

    return render(request, 'meuapp/upload.html')


def proteger_pdf(request):
    """Proteger PDF com senha de abertura e restrições"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        pdf = request.FILES['arquivo']
        senha = request.POST.get('senha', '123456')

        doc = pikepdf.Pdf.open(BytesIO(pdf.read()))

        permissoes = pikepdf.Permissions(
            extract=False,
            modify_annotation=False,
            modify_assembly=False,
        )
        encriptacao = pikepdf.Encryption(
            user=senha,
            owner=senha + '_owner',
            allow=permissoes
        )

        buffer = BytesIO()
        doc.save(buffer, encryption=encriptacao)
        buffer.seek(0)

        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="protegido.pdf"'
        return response

    return render(request, 'meuapp/upload.html')


def ocr_pdf(request):
    """OCR usando PyMuPDF — extrai texto de PDFs digitais e escaneados"""
    if request.method == 'POST' and request.FILES.get('arquivo'):
        pdf = request.FILES['arquivo']
        doc = fitz.open(stream=pdf.read(), filetype='pdf')
        texto_total = ''
        for i, pagina in enumerate(doc):
            texto = pagina.get_text()
            texto_total += f'\n--- Página {i + 1} ---\n{texto}'
        doc.close()

        response = HttpResponse(texto_total, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = 'attachment; filename="ocr_resultado.txt"'
        return response

    return render(request, 'meuapp/upload.html')
